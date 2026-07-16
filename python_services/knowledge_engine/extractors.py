"""
Knowledge Engine — Document Extractors
Extracts plain text + page/section metadata from PDF, DOCX, XLSX, CSV, ZIP.
Reuses pdfplumber (already a project dependency).
"""
from __future__ import annotations
import csv
import io
import logging
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

logger = logging.getLogger("cip.knowledge_engine.extractors")


@dataclass
class ExtractedPage:
    page_number:   Optional[int]
    text:          str
    section_label: Optional[str] = None


@dataclass
class ExtractionResult:
    pages:        list[ExtractedPage]
    full_text:    str
    page_count:   int
    extractor:    str
    warnings:     list[str]


# Section heading patterns common in airline operations manuals
SECTION_PATTERN = re.compile(
    r'^\s*(\d{1,2}(?:\.\d{1,2}){1,3})\s+[A-Z]', re.MULTILINE
)


def extract_document(file_path: str, file_type: str) -> ExtractionResult:
    """Dispatch to the correct extractor based on file type."""
    file_type = file_type.upper()
    if file_type == "PDF":
        return _extract_pdf(file_path)
    if file_type == "DOCX":
        return _extract_docx(file_path)
    if file_type == "XLSX":
        return _extract_xlsx(file_path)
    if file_type == "CSV":
        return _extract_csv(file_path)
    if file_type == "ZIP":
        return _extract_zip(file_path)
    raise ValueError(f"Unsupported file type for extraction: {file_type}")


def _extract_pdf(file_path: str) -> ExtractionResult:
    import pdfplumber

    pages: list[ExtractedPage] = []
    warnings: list[str] = []
    full_text_parts: list[str] = []
    current_section: Optional[str] = None

    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
                if not text.strip():
                    warnings.append(f"Page {i}: no extractable text")
                    continue

                match = SECTION_PATTERN.search(text)
                if match:
                    current_section = match.group(1)

                pages.append(ExtractedPage(
                    page_number=i,
                    text=text,
                    section_label=current_section,
                ))
                full_text_parts.append(text)
    except Exception as e:
        warnings.append(f"PDF extraction error: {e}")

    return ExtractionResult(
        pages=pages,
        full_text="\n\n".join(full_text_parts),
        page_count=len(pages),
        extractor="pdfplumber",
        warnings=warnings,
    )


def _extract_docx(file_path: str) -> ExtractionResult:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError("python-docx not installed")

    warnings: list[str] = []
    pages: list[ExtractedPage] = []
    current_section: Optional[str] = None
    full_text_parts: list[str] = []

    try:
        doc = DocxDocument(file_path)
        buffer: list[str] = []
        block_index = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = para.style.name.lower().startswith("heading")
            match = SECTION_PATTERN.match(text)
            if match:
                current_section = match.group(1)

            if is_heading and buffer:
                block_index += 1
                joined = "\n".join(buffer)
                pages.append(ExtractedPage(
                    page_number=block_index,
                    text=joined,
                    section_label=current_section,
                ))
                full_text_parts.append(joined)
                buffer = []

            buffer.append(text)

        if buffer:
            block_index += 1
            joined = "\n".join(buffer)
            pages.append(ExtractedPage(
                page_number=block_index,
                text=joined,
                section_label=current_section,
            ))
            full_text_parts.append(joined)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(c.text.strip() for c in row.cells))
            if rows:
                block_index += 1
                table_text = "\n".join(rows)
                pages.append(ExtractedPage(
                    page_number=block_index,
                    text=table_text,
                    section_label=current_section,
                ))
                full_text_parts.append(table_text)

    except Exception as e:
        warnings.append(f"DOCX extraction error: {e}")

    return ExtractionResult(
        pages=pages,
        full_text="\n\n".join(full_text_parts),
        page_count=len(pages),
        extractor="python-docx",
        warnings=warnings,
    )


def _extract_xlsx(file_path: str) -> ExtractionResult:
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("openpyxl not installed")

    warnings: list[str] = []
    pages: list[ExtractedPage] = []
    full_text_parts: list[str] = []

    try:
        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        for sheet_idx, sheet_name in enumerate(wb.sheetnames, start=1):
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                sheet_text = f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text)
                pages.append(ExtractedPage(
                    page_number=sheet_idx,
                    text=sheet_text,
                    section_label=sheet_name,
                ))
                full_text_parts.append(sheet_text)
    except Exception as e:
        warnings.append(f"XLSX extraction error: {e}")

    return ExtractionResult(
        pages=pages,
        full_text="\n\n".join(full_text_parts),
        page_count=len(pages),
        extractor="openpyxl",
        warnings=warnings,
    )


def _extract_csv(file_path: str) -> ExtractionResult:
    warnings: list[str] = []
    pages: list[ExtractedPage] = []
    text = ""

    try:
        with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows = [" | ".join(row) for row in reader if row]
        text = "\n".join(rows)
        pages.append(ExtractedPage(page_number=1, text=text))
    except Exception as e:
        warnings.append(f"CSV extraction error: {e}")

    return ExtractionResult(
        pages=pages,
        full_text=text,
        page_count=len(pages),
        extractor="csv",
        warnings=warnings,
    )


def _extract_zip(file_path: str) -> ExtractionResult:
    """
    Extract all supported files inside a ZIP and concatenate results.
    Each inner file becomes a labeled section.
    """
    warnings: list[str] = []
    pages: list[ExtractedPage] = []
    full_text_parts: list[str] = []
    page_counter = 0

    supported_ext = {".pdf": "PDF", ".docx": "DOCX", ".xlsx": "XLSX", ".csv": "CSV"}

    try:
        with zipfile.ZipFile(file_path) as zf:
            for name in zf.namelist():
                ext = Path(name).suffix.lower()
                if ext not in supported_ext:
                    continue
                try:
                    with zf.open(name) as inner_file:
                        data = inner_file.read()
                        tmp_path = f"/tmp/_zip_extract_{Path(name).name}"
                        with open(tmp_path, "wb") as tmp:
                            tmp.write(data)
                        inner_result = extract_document(tmp_path, supported_ext[ext])
                        for p in inner_result.pages:
                            page_counter += 1
                            pages.append(ExtractedPage(
                                page_number=page_counter,
                                text=p.text,
                                section_label=f"{name} :: {p.section_label or ''}",
                            ))
                        full_text_parts.append(f"[File: {name}]\n{inner_result.full_text}")
                        warnings.extend(
                            f"{name}: {w}" for w in inner_result.warnings)
                except Exception as e:
                    warnings.append(f"Failed to extract {name}: {e}")
    except Exception as e:
        warnings.append(f"ZIP extraction error: {e}")

    return ExtractionResult(
        pages=pages,
        full_text="\n\n".join(full_text_parts),
        page_count=len(pages),
        extractor="zipfile",
        warnings=warnings,
    )
